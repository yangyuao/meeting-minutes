from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm

logger = logging.getLogger(__name__)

CHINESE_NUMBERS = [
    "一",
    "二",
    "三",
    "四",
    "五",
    "六",
    "七",
    "八",
    "九",
    "十",
    "十一",
    "十二",
]


def clean_bold(text: str) -> str:
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text)


def extract_section(text: str, section_title: str) -> str:
    if section_title == "会议主题":
        escaped_title = re.escape(section_title)
        pattern = rf"^# {escaped_title}\s*\n(.*?)(?=^# |\Z)"
        match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
        if match:
            content = match.group(1).strip()
            lines = content.split("\n")
            for line in lines:
                if line.strip().startswith("- "):
                    return line.strip()[2:]
            return content
        return ""

    escaped_title = re.escape(section_title)
    pattern = rf"^## {escaped_title}\s*\n(.*?)(?=^## |\Z)"
    match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
    return match.group(1).strip() if match else ""


def fill_cell_with_formatting(cell, content: str) -> None:
    cell._element.clear_content()
    if not content:
        cell.text = "（无内容）"
        return

    lines = content.split("\n")
    h3_counter = 0
    list_counter = 0

    for line in lines:
        stripped_line = line.rstrip()
        if not stripped_line:
            paragraph = cell.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.75)
            continue

        if line.startswith("### "):
            h3_counter += 1
            list_counter = 0
            title_text = line[4:].strip()
            prefix = (
                f"{CHINESE_NUMBERS[h3_counter - 1]}、"
                if h3_counter <= len(CHINESE_NUMBERS)
                else f"{h3_counter}、"
            )
            paragraph = cell.add_paragraph(prefix + title_text)
            paragraph.paragraph_format.left_indent = Cm(0)
            continue

        if line.lstrip().startswith("- "):
            list_counter += 1
            item_text = line.lstrip()[2:].strip()
            item_text = clean_bold(item_text)
            numbered_text = f"{list_counter}. {item_text}"
            paragraph = cell.add_paragraph(numbered_text)
            paragraph.paragraph_format.left_indent = Cm(0.75)
            continue

        clean_line = clean_bold(stripped_line)
        paragraph = cell.add_paragraph(clean_line)
        paragraph.paragraph_format.left_indent = Cm(0.75)


def generate_docx_file(
    markdown_text: str,
    template_path: Path,
    output_path: Path,
) -> Path:
    if not template_path.exists():
        raise FileNotFoundError(f"Template file '{template_path}' not found")

    data = {
        "会议主题": extract_section(markdown_text, "会议主题"),
        "会议要点": extract_section(markdown_text, "讨论议题"),
        "待办事项": extract_section(markdown_text, "决定和行动计划"),
    }

    doc = Document(str(template_path))
    filled = 0

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            key = row.cells[0].text.strip().rstrip("：").rstrip(":").strip()
            cell = row.cells[1]

            if key in data:
                fill_cell_with_formatting(cell, data[key])
                filled += 1

    doc.save(str(output_path))
    logger.info("DOCX generated with %s sections", filled)
    return output_path
