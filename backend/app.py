"""
会议纪要自动生成系统后端服务

本模块实现了基于FastAPI的Web服务，提供以下核心功能：
1. 语音转文字（使用funasr模型）
2. 文本清洗与预处理
3. 调用大语言模型生成会议纪要
"""

import re
import requests
from fastapi import FastAPI, UploadFile, File, Request, HTTPException, Form
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import json
import os
import torch
from datetime import datetime
from typing import Generator
from funasr import AutoModel
from datetime import timedelta
import uvicorn
from dotenv import load_dotenv
import logging
import sys

# 数据库相关导入
import pymysql
from pymysql import Error
from contextlib import contextmanager

# 在文件顶部添加新的导入
from docx import Document
from docx.shared import Cm

# 配置日志，确保与 Uvicorn 兼容
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    stream=sys.stdout,  # 明确指定输出流
    force=True,  # 覆盖任何现有的基本配置
)

logger = logging.getLogger(__name__)

# 加载环境变量
load_dotenv()

# 创建FastAPI应用实例
app = FastAPI()
# 添加CORS中间件以支持跨域请求
app.add_middleware(
    CORSMiddleware,
    allow_origins=os.getenv("CORS_ORIGINS", "http://localhost:5173").split(","),
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# API端点配置
OLLAMA_API = os.getenv("OLLAMA_API", "http://localhost:11434/api/generate")
LLM_API = os.getenv("LLM_API", "http://10.8.3.173:1025/v1/api/generate")
VLLM_API = os.getenv("VLLM_API", "http://10.8.2.21:8999/v1/chat/completions")
PRODUCTION_IP = os.getenv("PRODUCTION_IP", "10.6.2.178")

# 数据库配置
DB_CONFIG = {
    "host": os.getenv("DB_HOST", "localhost"),
    "port": int(os.getenv("DB_PORT", 3306)),
    "user": os.getenv("DB_USER", "root"),
    "password": os.getenv("DB_PASSWORD", ""),
    "database": os.getenv("DB_NAME", "meeting_minutes"),
    "charset": "utf8mb4",
}

# 提示词模板
PROMPT_TEMPLATE = os.getenv(
    "PROMPT_TEMPLATE",
    """## Background
语音记录会议讨论信息，现在可以方便地转成文字。但这些碎片信息，如何方便整理成没有口语、逻辑清晰、内容明确的会议纪要。

## Goals
- 保证会议内容被全面地记录、准确地表述。准确记录会议的各个方面，包括议题、讨论、决定和行动计划。
- 保证语言通畅，易于理解，使每个参会人员都能明确理解会议内容框架和结论。
- 生成结构化、准确、不臆想、可直接用于对上级汇报的正式会议纪要。

## Workflow
- 输入：用户提供一段原始会议记录文本；用户提供会议讨论的基本信息。
- 整理：遵循以下框架来整理用户提供的会议信息，每个步骤后都会进行数据校验确保信息准确性
- 输出：输出整理后的结构清晰，描述完整的会议纪要。

## Output format requirements
- 必须严格使用markdown格式
- 层级标题使用# ## ### ####符号，不得使用数字编号
- 列表项使用-符号，不得使用数字编号
- 不得包含任何阿拉伯数字、中文数字、特殊符号
- 段落之间保持合理间距

## Tone
- 专业：使用专业术语和格式
- 简洁：信息要点明确，不做多余的解释
- 严格遵循格式规范：绝对禁止使用任何数字编号

## Constraints
- 整理会议纪要过程中，需严格遵守信息准确性，不对用户提供的信息做扩写，不得凭空编造未出现的信息。
- 仅做信息整理，将一些明显的病句做微调。
- 不得根据推测生成决定、任务或分工。
- 决定和行动计划与下一步打算是一个意思，要避免重复，只写决定和行动计划。
- 如果发现无法判断某项内容，请输出：未提及或信息不足
- 如果发现内容与议题不相关，必须归到其他事项。
- 当用户提供了议题时，必须严格按照议题组织会议纪要；若未提供议题，你需要从内容中自动提取议题并组织结构。
- 绝对禁止在任何部分使用数字编号，包括但不限于：标题编号、列表编号、序号等
- 所有列表必须使用-符号，所有标题必须使用#符号

## Skills
- 文字处理：具备优秀的文字组织和编辑能力。去除所有口水话、废话、重复内容。避免冗长叙述、重复内容、口语化表达。
- 格式控制：严格遵守markdown格式规范
- 内容层级：会议中真实的内容只放在以 ‘-’ 开始的列表项里，不要相互之间有层级，-的上级只能是 ###的标题项,  一切列表项都以‘-’开头

## Value
- 准确性：确保记录的信息无误
- 格式一致性：确保输出格式完全符合要求

## Output example

# 会议主题
- DRC业务讨论
## 讨论议题
### 工作汇报
- 发言人一汇报了本周工作：基于通用规则构建了业务建模框架，初步测试生成规则的准确性，发现模型在少量规则生成时表现良好，但需优化处理复杂规则的逻辑。
- 提出需进一步优化业务建模工具，支持历史案例的分类与规则提取，并计划下周推进业务建模逻辑设计。

### 待优化项
- 发言人三强调业务建模需总结历史案例，通过工具解析文档，提取规则并关联到逻辑结构，形成可复用的规则库。
- 讨论模型生成规则的准确性问题：当前模型在处理类似规则时易混淆，需分批次生成并结合示例优化结果。

## 决定和行动计划
### 模型测试
- 使用千问三模型进行规则生成测试，验证其在分批次生成场景下的准确性。
- 优化业务建模工具，支持规则分类与历史案例关联。

### 知识库更新
- 完善知识库结构，补充文档更新历史记录，确保模型可追溯文档变更。
- 制定知识库更新规范，要求上传文档时标注版本信息。
 
""",
)

# 全局变量用于存储模型实例
_model_instance = None


def get_model():
    """
    获取模型实例，利用模块级变量实现模型缓存
    在Uvicorn热重载时避免重复加载模型
    """
    global _model_instance
    if _model_instance is None:
        logger.info("正在加载语音识别模型...")
        _model_instance = AutoModel(
            model=os.getenv(
                "ASR_MODEL",
                "/root/.cache/modelscope/hub/models/iic/speech_seaco_paraformer_large_asr_nat-zh-cn-16k-common-vocab8404-pytorch",
            ),
            vad_model=os.getenv(
                "VAD_MODEL",
                "/root/.cache/modelscope/hub/models/iic/speech_fsmn_vad_zh-cn-16k-common-pytorch",
            ),
            punc_model=os.getenv(
                "PUNC_MODEL",
                "/root/.cache/modelscope/hub/models/iic/punc_ct-transformer_cn-en-common-vocab471067-large",
            ),
            spk_model=os.getenv(
                "SPK_MODEL",
                "/root/.cache/modelscope/hub/models/iic/speech_campplus_sv_zh-cn_16k-common",
            ),
            disable_update=True,
            device=os.getenv("DEVICE", "cuda:0"),
        )
        logger.info("语音识别模型加载完成")
    else:
        logger.info("使用已缓存的模型实例")
    return _model_instance


# 初始化模型
model = get_model()


@contextmanager
def get_db_connection():
    """
    数据库连接上下文管理器
    """
    connection = None
    try:
        connection = pymysql.connect(**DB_CONFIG)
        yield connection
    except Error as e:
        logger.error(f"数据库连接错误: {e}")
        if connection:
            connection.rollback()
        raise
    finally:
        if connection:
            connection.close()


def init_database():
    """
    初始化数据库表
    """
    try:
        with get_db_connection() as connection:
            with connection.cursor() as cursor:
                # 创建会议记录表
                create_table_query = """
                CREATE TABLE IF NOT EXISTS meeting_minutes_records (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    client_ip VARCHAR(45) NOT NULL,
                    username VARCHAR(255),
                    original_text LONGTEXT,
                    final_prompt LONGTEXT,
                    generated_summary LONGTEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_client_ip (client_ip),
                    INDEX idx_username (username),
                    INDEX idx_created_at (created_at)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
                """
                cursor.execute(create_table_query)

                # 创建转录文件记录表
                create_transcript_table_query = """
                CREATE TABLE IF NOT EXISTS meeting_minutes_transcript_files (
                    id INT AUTO_INCREMENT PRIMARY KEY,
                    client_ip VARCHAR(45) NOT NULL,
                    username VARCHAR(255),
                    audio_file_path VARCHAR(512),
                    transcript_file_path VARCHAR(512),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    INDEX idx_client_ip (client_ip),
                    INDEX idx_username (username),
                    INDEX idx_created_at (created_at)
                ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci
                """
                cursor.execute(create_transcript_table_query)

                connection.commit()
                logger.info("数据库表初始化完成")
    except Exception as e:
        logger.error(f"数据库初始化失败: {e}")


def save_meeting_record(
    client_ip: str,
    username: str,
    original_text: str,
    final_prompt: str,
    generated_summary: str,
):
    """
    保存会议记录到数据库

    Args:
        client_ip (str): 客户端IP地址
        username (str): 用户名
        original_text (str): 原始文字稿
        final_prompt (str): 最终提示词
        generated_summary (str): 生成的会议纪要
    """
    try:
        with get_db_connection() as connection:
            with connection.cursor() as cursor:
                insert_query = """
                INSERT INTO meeting_minutes_records (client_ip, username, original_text, final_prompt, generated_summary)
                VALUES (%s, %s, %s, %s, %s)
                """
                cursor.execute(
                    insert_query,
                    (
                        client_ip,
                        username,
                        original_text,
                        final_prompt,
                        generated_summary,
                    ),
                )
                connection.commit()
                logger.info(f"会议记录已保存到数据库，ID: {cursor.lastrowid}")
    except Exception as e:
        logger.error(f"保存会议记录失败: {e}")


def save_transcript_record(
    client_ip: str,
    username: str,
    audio_file_path: str,
    transcript_file_path: str,
):
    """
    保存转录文件记录到数据库

    Args:
        client_ip (str): 客户端IP地址
        username (str): 用户名
        audio_file_path (str): 音频文件路径
        transcript_file_path (str): 转录文本文件路径
    """
    try:
        with get_db_connection() as connection:
            with connection.cursor() as cursor:
                insert_query = """
                INSERT INTO meeting_minutes_transcript_files (client_ip, username, audio_file_path, transcript_file_path)
                VALUES (%s, %s, %s, %s)
                """
                cursor.execute(
                    insert_query,
                    (
                        client_ip,
                        username,
                        audio_file_path,
                        transcript_file_path,
                    ),
                )
                connection.commit()
                logger.info(f"转录文件记录已保存到数据库，ID: {cursor.lastrowid}")
    except Exception as e:
        logger.error(f"保存转录文件记录失败: {e}")


def get_client_ip(request: Request) -> str:
    """
    获取客户端真实IP地址（支持代理）

    Args:
        request (Request): HTTP请求对象

    Returns:
        str: 客户端真实IP地址
    """
    # 获取客户端真实IP地址（支持代理）
    forwarded_for = request.headers.get("x-forwarded-for")
    if forwarded_for:
        # X-Forwarded-For 可能包含多个IP，第一个是原始客户端IP
        client_ip = forwarded_for.split(",")[0].strip()
    else:
        # 如果没有X-Forwarded-For头，则尝试X-Real-IP
        real_ip = request.headers.get("x-real-ip")
        if real_ip:
            client_ip = real_ip
        else:
            # 最后回退到直接连接的客户端IP
            client_ip = request.client.host

    return client_ip


def clean_speech(text):
    """
    清理语音转文字后的文本，去除多余空格和口语化词汇

    该函数会对输入文本进行以下处理：
    1. 合并多个空格为单个空格
    2. 移除常见的口语化 filler 词汇（如"嗯"、"啊"等）
    3. 再次合并空格并去除首尾空白

    Args:
        text (str): 需要清理的原始文本

    Returns:
        str: 清理后的文本
    """
    # 合并多个空格为单个空格并去除首尾空白
    text = re.sub(r"\s+", " ", text).strip()
    # 定义需要移除的口语化词汇
    noise_words = ["嗯", "啊", "哦", "呃", "哈", "这个", "那个", "然后", "就是", "的话"]
    # 移除文本中的口语化词汇
    for word in noise_words:
        text = re.sub(r"\b" + re.escape(word) + r"\b", "", text)
    # 再次合并空格并去除首尾空白
    text = re.sub(r"\s+", " ", text).strip()
    return text


def generate_ollama_stream(final_prompt: str) -> Generator[str, None, None]:
    """
    使用Ollama模型流式生成文本

    Args:
        final_prompt (str): 发送给模型的完整提示词

    Yields:
        str: 模型生成的文本片段
    """
    # 构建请求载荷
    payload = {
        "model": os.getenv("OLLAMA_MODEL", "qwen3:4b"),
        "prompt": final_prompt,
        "stream": True,
    }

    try:
        # 发送POST请求到Ollama API
        with requests.post(OLLAMA_API, json=payload, stream=True) as r:
            # 检查响应状态码
            if r.status_code != 200:
                logger.error("无法连接到 Ollama 模型，状态码: %d", r.status_code)
                yield "[错误] 无法连接到 Ollama 模型。"
                return

            # 处理流式响应数据
            for chunk in r.iter_content(chunk_size=None):
                if chunk:
                    try:
                        text = chunk.decode("utf-8").strip()
                        data = json.loads(text)
                        response = data.get("response", "")
                        if response:
                            yield response
                    except json.JSONDecodeError:
                        continue
    except Exception as e:
        logger.error("请求 Ollama 模型时出错: %s", str(e))
        yield f"[错误] 请求 Ollama 模型时出错: {str(e)}"


@app.get("/test-vllm")
async def test_vllm():
    """
    测试vLLM模型的API端点

    Returns:
        StreamingResponse: 流式响应对象
    """
    return StreamingResponse(
        generate_vllm_stream(system_prompt=PROMPT_TEMPLATE, user_prompt="请介绍你自己"),
        media_type="text/plain",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


def generate_vllm_stream(
    system_prompt: str, user_prompt: str
) -> Generator[str, None, None]:
    """
    使用vLLM模型流式生成文本

    Args:
        final_prompt (str): 发送给模型的完整提示词

    Yields:
        str: 模型生成的文本片段
    """
    # 构建请求载荷，使用PROMPT_TEMPLATE作为系统提示词但不含transcript内容
    system_prompt = system_prompt if len(system_prompt) > 0 else PROMPT_TEMPLATE

    payload = {
        "model": os.getenv("VLLM_MODEL", "/home/ubuntu/models/Qwen3-8B"),
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.00,
        "stream": True,
    }

    try:
        # 发送POST请求到vLLM API
        with requests.post(VLLM_API, json=payload, stream=True, timeout=10) as r:
            # 检查响应状态码
            if r.status_code != 200:
                logger.error("无法连接到 vLLM 模型，状态码: %d", r.status_code)
                logger.error("响应内容: %s", r.text)
                yield f"响应内容: {str(r.text)}"
                return

            # 处理流式响应数据
            for chunk in r.iter_content(chunk_size=None):
                if chunk:
                    try:
                        text = chunk.decode("utf-8").strip()
                        if text.startswith("data:"):
                            data_str = text[5:].strip()
                            if data_str == "[DONE]":
                                break
                            data = json.loads(data_str)
                            content = (
                                data.get("choices", [{}])[0]
                                .get("delta", {})
                                .get("content", "")
                            )
                            if content:
                                yield content
                    except json.JSONDecodeError:
                        continue
    except Exception as e:
        logger.error("请求 vLLM 模型时出错: %s", str(e))
        yield f"[错误] 请求 vLLM 模型时出错: {str(e)}"


def unified_streamer(
    system_prompt: str, user_prompt: str, is_production: bool
) -> Generator[str, None, None]:
    """
    统一流式响应处理器，根据环境选择合适的模型

    Args:
        final_prompt (str): 发送给模型的完整提示词
        is_production (bool): 是否为生产环境

    Yields:
        str: 模型生成的文本片段
    """
    if not is_production:
        yield from generate_ollama_stream(final_prompt=system_prompt + user_prompt)
    else:
        yield from generate_vllm_stream(
            system_prompt=system_prompt, user_prompt=user_prompt
        )


@app.get("/")
async def read_root():
    """
    根路径，用于检查服务是否运行正常

    Returns:
        dict: 包含服务状态消息的字典
    """
    logger.info("访问根路径，服务运行正常")
    return {"message": "Backend is running"}


@app.get("/prompt/default")
async def get_default_prompt():
    """
    获取默认提示词模板

    Returns:
        dict: 包含默认提示词的字典
    """
    logger.info("获取默认提示词模板")
    return {"prompt": PROMPT_TEMPLATE}


def process_transcript(input_file, output_file):
    """
    处理语音识别结果，按说话人分组句子

    Args:
        input_file (str): 输入JSON文件路径
        output_file (str): 输出JSON文件路径
    """
    logger.debug("开始处理转录文件: %s", input_file)
    # 读取输入文件
    with open(input_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    # 移除时间戳信息
    if data[0]["timestamp"]:
        data[0].pop("timestamp")

    # 初始化处理结果列表和当前说话人相关变量
    processed = []
    current_spk = None
    current_batch = []

    # 遍历句子信息
    for item in data[0]["sentence_info"]:
        try:
            if not isinstance(item, dict):
                continue

            # 提取句子相关信息
            spk = item.get("spk")
            text = str(item.get("text", ""))
            start = float(item.get("start", 0))
            end = float(item.get("end", 0))

            # 按说话人分组处理句子
            if spk == current_spk:
                current_batch.append({"text": text, "start": start, "end": end})
            else:
                if current_batch:
                    processed.append(
                        {
                            "text": " ".join([s["text"] for s in current_batch]),
                            "spk": current_spk,
                            "start": min(s["start"] for s in current_batch),
                            "end": max(s["end"] for s in current_batch),
                        }
                    )
                current_batch = [{"text": text, "start": start, "end": end}]
                current_spk = spk

        except Exception as e:
            logger.error("处理转录数据时发生异常: %s", str(e))
            continue

    # 处理最后一批数据
    if current_batch and current_batch[0].get("text"):
        processed.append(
            {
                "text": " ".join([s["text"] for s in current_batch]),
                "spk": current_spk,
                "start": min(s["start"] for s in current_batch),
                "end": max(s["end"] for s in current_batch),
            }
        )
    data[0]["sentence_info"] = processed

    # 将处理结果写入输出文件
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    logger.debug("转录处理完成，输出文件: %s", output_file)


def ms_to_timestamp(milliseconds):
    """
    将毫秒转换为HH:MM:SS格式的时间戳

    Args:
        milliseconds (float): 毫秒数

    Returns:
        str: 格式化后的时间字符串
    """
    seconds = milliseconds / 1000
    td = timedelta(seconds=seconds)
    return str(td).split(".")[0]


def process_json(input_file, output_file):
    """
    将JSON格式的语音识别结果转换为文本格式

    Args:
        input_file (str): 输入JSON文件路径
        output_file (str): 输出文本文件路径
    """
    logger.debug("开始处理JSON文件: %s", input_file)
    # 读取输入文件
    with open(input_file, "r", encoding="utf-8") as f:
        data = json.load(f)

    # 写入输出文件
    with open(output_file, "w", encoding="utf-8") as f:
        for item in data:
            for sentence in item["sentence_info"]:
                # 提取句子相关信息
                speaker = f"发言人{sentence['spk']}"
                text = "".join(sentence["text"].split(" "))

                # 写入格式化的文本
                # start_time = ms_to_timestamp(sentence["start"])
                # end_time = ms_to_timestamp(sentence["end"])
                # f.write(f"{start_time}-{end_time}\n")
                f.write(f"{speaker}: {text}\n")
                # f.write("\n")
    logger.debug("JSON处理完成，输出文件: %s", output_file)


@app.post("/get_transcript")
async def get_transcript(
    request: Request, file: UploadFile = File(...), username: str = Form("unknown")
):
    """
    上传音频文件并获取语音识别结果

    Args:
        request (Request): HTTP请求对象，用于获取客户端IP
        file (UploadFile): 上传的音频文件
        username (str): 用户名，默认为"unknown"

    Returns:
        dict: 包含语音识别结果或错误信息的字典
    """
    # 获取客户端真实IP地址（支持代理）
    client_ip = get_client_ip(request)
    logger.info("进入 get_transcript 接口，客户端IP: %s", client_ip)

    # 创建基于用户名、IP和时间戳的目录结构
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    ip_dir = client_ip.replace(".", "_").replace(":", "_")  # 处理IPv6地址
    # 修改存储路径到当前目录下的audio_files文件夹
    save_dir = f"./audio_files/{username}/{ip_dir}/{timestamp}"
    os.makedirs(save_dir, exist_ok=True)

    # 保存上传文件到基于用户名、IP和时间戳的目录中
    temp_path = f"{save_dir}/{file.filename}"
    with open(temp_path, "wb") as buffer:
        content = await file.read()
        logger.info("文件大小: %d 字节", len(content))
        buffer.write(content)

    # 记录文件保存位置
    logger.info("文件已保存到: %s", temp_path)

    # 生成临时文件名用于处理过程
    temp_json = f"{save_dir}/temp_{timestamp}.json"
    processed_json = f"{save_dir}/processed_{timestamp}.json"
    output_txt = f"{save_dir}/output_{timestamp}.txt"

    try:
        logger.info("开始语音识别...")
        # 使用funasr模型进行语音识别
        res = model.generate(
            input=temp_path,
            batch_size_s=64,
            hotword="./hotwords.txt",
        )
        torch.set_num_threads(4)
        # 保存识别结果
        with open(temp_json, "w", encoding="utf-8") as f:
            json.dump(res, f, ensure_ascii=False, indent=2)

        # 处理识别结果
        process_transcript(temp_json, processed_json)
        process_json(processed_json, output_txt)

        # 读取最终文本结果
        with open(output_txt, "r") as buffer:
            full_text = buffer.read()

        logger.info("语音识别完成")
        logger.debug("语音识别结果: %s", full_text)

        # 保存转录记录到数据库
        save_transcript_record(client_ip, username, temp_path, output_txt)

        # 注意：我们不再删除保存的文件，以便永久保留
        # 只删除临时的JSON文件
        os.remove(temp_json)
        os.remove(processed_json)

        return {
            "transcript": full_text,
            "audio_file_path": temp_path,
            "transcript_file_path": output_txt,
        }

    except Exception as e:
        logger.error("语音识别失败: %s", str(e))
        # 即使处理失败，我们也保留原始音频文件用于调试
        return {"error": str(e)}


@app.post("/generate_summary")
async def generate_summary(request: Request):
    """
    基于语音识别文本生成会议纪要

    Args:
        request (Request): 包含文本和提示词的HTTP请求

    Returns:
        StreamingResponse: 流式响应对象，包含生成的会议纪要
    """
    # 获取客户端真实IP地址
    client_ip = get_client_ip(request)
    logger.info("进入 generate_summary 接口，客户端IP: %s", client_ip)

    generated_summary = ""

    try:
        # 解析请求数据
        request_data = await request.json()

        # 提取文本和提示词
        full_text: str = request_data.get("transcript", "")
        prompt: str = request_data.get("prompt", "")

        # 清洗文本
        cleaned_text = clean_speech(full_text)
        logger.info("文本清洗完成")
        logger.debug("文本清洗结果: %s", cleaned_text)

        # 构建最终提示词
        final_prompt = (
            (prompt.strip() + "\n\n" + cleaned_text)
            if prompt.strip()
            else PROMPT_TEMPLATE + "\n\n" + cleaned_text
        )
        logger.info("构建最终提示词完成")

        # 判断是否为生产环境
        is_production = client_ip == PRODUCTION_IP
        logger.info("正在返回 StreamingResponse (生产环境: %s)", is_production)

        # 定义一个包装生成器来捕获生成的内容
        def summary_streamer():
            nonlocal generated_summary
            for chunk in unified_streamer(
                system_prompt=prompt, user_prompt=cleaned_text, is_production=True
            ):
                generated_summary += chunk
                yield chunk

            # 在流式传输完成后保存记录到数据库
            try:
                # 获取额外参数
                username = request_data.get("username", "unknown")
                # 不再记录音频文件和transcript文件的路径
                save_meeting_record(
                    client_ip, username, full_text, final_prompt, generated_summary
                )
                logger.info("会议记录已保存到数据库")
            except Exception as e:
                logger.error(f"保存会议记录到数据库失败: {e}")

        # 返回流式响应
        return StreamingResponse(
            summary_streamer(),
            media_type="text/plain",
            headers={
                "Content-Type": "text/plain; charset=utf-8",
                "Cache-Control": "no-cache",
                "X-Accel-Buffering": "no",  # 关键！防止 Nginx 缓冲流式内容
            },
        )

    except Exception as e:
        logger.error("请求失败: %s", str(e))
        raise HTTPException(
            status_code=500,
            detail=f"请求失败，请查看控制台错误信息。详细错误: {str(e)}",
        )


@app.post("/upload")
async def upload_audio(
    request: Request, file: UploadFile = File(...), username: str = Form("unknown")
):
    """
    上传音频文件并保存到服务器

    Args:
        request (Request): HTTP请求对象，用于获取客户端IP
        file (UploadFile): 上传的音频文件
        username (str): 用户名，默认为"unknown"

    Returns:
        dict: 包含文件保存路径的信息
    """
    # 获取客户端真实IP地址（支持代理）
    client_ip = get_client_ip(request)
    logger.info("进入 upload 接口，客户端IP: %s", client_ip)

    # 创建基于用户名、IP和时间戳的目录结构
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    ip_dir = client_ip.replace(".", "_").replace(":", "_")  # 处理IPv6地址
    # 修改存储路径到当前目录下的audio_files文件夹
    save_dir = f"./audio_files/{username}/{ip_dir}/{timestamp}"
    os.makedirs(save_dir, exist_ok=True)

    # 保存上传文件到基于用户名、IP和时间戳的目录中
    file_path = f"{save_dir}/{file.filename}"
    with open(file_path, "wb") as buffer:
        content = await file.read()
        logger.info("文件大小: %d 字节", len(content))
        buffer.write(content)

    # 记录文件保存位置
    logger.info("文件已保存到: %s", file_path)

    # 保存转录记录到数据库（这里保存的是原始音频文件的记录）
    save_transcript_record(client_ip, username, file_path, "")

    return {
        "message": "文件上传成功",
        "file_path": file_path,
        "client_ip": client_ip,
        "username": username,
    }


CHINESE_NUMBERS = [
    "一",
    "二",
    "三",
    "四",
    "五",
    "六",
    "七",
    "八",
    "九",
    "十",
    "十一",
    "十二",
]


def clean_bold(text):
    """移除 **text** 中的星号，保留内容"""
    return re.sub(r"\*\*(.*?)\*\*", r"\1", text)


def extract_section(text, section_title):
    """Extract content under a specific heading"""
    # Handle the special case of "会议主题" which uses # instead of ##
    if section_title == "会议主题":
        escaped_title = re.escape(section_title)
        pattern = rf"^# {escaped_title}\s*\n(.*?)(?=^# |\Z)"
        match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
        if match:
            content = match.group(1).strip()
            # Extract the actual topic from the list item
            lines = content.split("\n")
            for line in lines:
                if line.strip().startswith("- "):
                    return line.strip()[2:]  # Remove "- " prefix
            return content
        return ""
    else:
        # Original logic for other sections with ## headers
        escaped_title = re.escape(section_title)
        pattern = rf"^## {escaped_title}\s*\n(.*?)(?=^## |\Z)"
        match = re.search(pattern, text, re.DOTALL | re.MULTILINE)
        return match.group(1).strip() if match else ""


def fill_cell_with_formatting(cell, content):
    cell._element.clear_content()
    if not content:
        cell.text = "（无内容）"
        return

    lines = content.split("\n")
    h3_counter = 0
    list_counter = 0  # 每个 ### 标题下的列表计数器

    for line in lines:
        stripped_line = line.rstrip()
        if not stripped_line:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            continue

        # === 遇到 ### 标题：重置列表计数器 ===
        if line.startswith("### "):
            h3_counter += 1
            list_counter = 0  # 重置编号
            title_text = line[4:].strip()
            prefix = (
                f"{CHINESE_NUMBERS[h3_counter - 1]}、"
                if h3_counter <= len(CHINESE_NUMBERS)
                else f"{h3_counter}、"
            )
            p = cell.add_paragraph(prefix + title_text)
            p.paragraph_format.left_indent = Cm(0)
            continue

        # === 遇到 - 列表项：转为 1. 2. 3. ===
        if line.lstrip().startswith("- "):
            list_counter += 1
            item_text = line.lstrip()[2:].strip()
            item_text = clean_bold(item_text)
            numbered_text = f"{list_counter}. {item_text}"
            p = cell.add_paragraph(numbered_text)
            p.paragraph_format.left_indent = Cm(0.75)
            continue

        # === 其他普通行（如"讨论："本身、或子说明）===
        clean_line = clean_bold(stripped_line)
        p = cell.add_paragraph(clean_line)
        p.paragraph_format.left_indent = Cm(0.75)


@app.post("/generate_docx")
async def generate_docx(request: Request):
    try:
        request_data = await request.json()
        markdown_text: str = request_data.get("markdown", "")

        if not markdown_text:
            raise HTTPException(status_code=400, detail="Missing markdown content")

        # 提取各部分内容
        data = {
            "会议主题": extract_section(markdown_text, "会议主题"),
            "会议要点": extract_section(markdown_text, "讨论议题"),
            "待办事项": extract_section(markdown_text, "决定和行动计划"),
        }

        # 创建临时文件路径
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        template_path = os.getenv("PPT_TEMPLATE", "./会议纪要模板.docx")
        output_path = f"/tmp/meeting_minutes_{timestamp}.docx"

        # 检查模板文件是否存在
        if not os.path.exists(template_path):
            raise HTTPException(
                status_code=500, detail=f"Template file '{template_path}' not found"
            )

        # 加载模板并填充内容
        doc = Document(template_path)
        filled = 0

        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) < 2:
                    continue
                key = row.cells[0].text.strip().rstrip("：").rstrip(":").strip()
                cell = row.cells[1]

                if key in data:
                    fill_cell_with_formatting(cell, data[key])
                    filled += 1

        # 保存文件
        doc.save(output_path)

        # 返回文件
        return FileResponse(
            path=output_path,
            filename=f"meeting_minutes_{timestamp}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=meeting_minutes_{timestamp}.docx"
            },
        )

    except Exception as e:
        print(f"Failed to generate DOCX: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate DOCX. Error: {str(e)}",
        )


# 应用程序入口点
if __name__ == "__main__":
    """
    启动FastAPI应用服务器
    """
    # 初始化数据库
    init_database()

    logger.info("启动FastAPI应用服务器")
    uvicorn.run(
        "app:app",
        host=os.getenv("HOST", "0.0.0.0"),
        port=int(os.getenv("PORT", 8002)),
        reload=os.getenv("DEBUG", "False").lower() == "true",
    )
